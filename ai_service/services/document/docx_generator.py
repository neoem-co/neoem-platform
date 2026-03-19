"""
DOCX generator for finalised contracts.

The layout is code-driven rather than relying on Word Title/Heading theme
styles, so the generated document keeps a more stable Thai legal format.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Optional

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from models.contract_draft import ContractArticle, PartyInfo
from services.document.layout_profile import (
    ThaiLegalLayoutProfile,
    get_default_thai_legal_layout_profile,
)

logger = logging.getLogger(__name__)

_NUMBERED_PREFIX = re.compile(
    r"^(?:ข้อ\s*\d+|"
    r"\d+(?:\.\d+)*[\.)]?|"
    r"\([ก-๙a-zA-Z0-9]+\)|"
    r"[ก-๙][\.)])"
)


def _is_structured_line(text: str) -> bool:
    return bool(_NUMBERED_PREFIX.match(text.strip()))


def _iter_paragraphs(text: str) -> list[str]:
    """Normalize text while preserving legal list/new-line structure."""
    if not text:
        return []

    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks = re.split(r"\n\s*\n+", normalized)
    paragraphs: list[str] = []

    for block in blocks:
        lines = [part.strip() for part in block.split("\n") if part.strip()]
        if not lines:
            continue

        buffer: list[str] = []
        for line in lines:
            line = re.sub(r"[ \t]{2,}", " ", line).strip()
            if not line:
                continue

            if _is_structured_line(line):
                if buffer:
                    paragraphs.append(" ".join(buffer).strip())
                    buffer = []
                paragraphs.append(line)
            else:
                buffer.append(line)

        if buffer:
            paragraphs.append(" ".join(buffer).strip())

    return paragraphs


def _resolve_font_name(profile: ThaiLegalLayoutProfile) -> str:
    return profile.docx_font_name or profile.docx_font_fallback


def _set_run_font(run, size: int, font_name: str, bold: bool = False):
    """Set a stable Thai font on a python-docx Run, including East Asia font."""
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(size)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def _resolve_paragraph_alignment(text: str):
    """Prefer Thai justify for real paragraphs, fallback for short/structured lines."""
    thai_justify = getattr(WD_ALIGN_PARAGRAPH, "THAI_JUSTIFY", WD_ALIGN_PARAGRAPH.JUSTIFY)
    stripped = text.strip()
    if len(stripped) < 32:
        return WD_ALIGN_PARAGRAPH.LEFT
    if _is_structured_line(stripped) and len(stripped) < 56:
        return WD_ALIGN_PARAGRAPH.LEFT
    return thai_justify


def _set_table_borderless(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_cell_width(cell, width_mm: float) -> None:
    cell.width = Mm(width_mm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_mm * 56.7)))


def _party_display_name(party: PartyInfo) -> str:
    return (party.company or party.name or "........................................").strip()


def _party_role_label(party: PartyInfo) -> str:
    role = (party.role or "").lower()
    mapping = {
        "buyer": "ผู้ว่าจ้าง",
        "client": "ผู้ว่าจ้าง",
        "seller": "ผู้รับจ้าง",
        "vendor": "ผู้รับจ้าง",
        "factory": "ผู้รับจ้าง",
    }
    return mapping.get(role, "คู่สัญญา")


def _new_paragraph(doc: DocxDocument):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    return para


def _write_body_paragraph(
    doc: DocxDocument,
    text: str,
    profile: ThaiLegalLayoutProfile,
    font_name: str,
    *,
    preamble: bool = False,
) -> None:
    para = _new_paragraph(doc)
    para.alignment = _resolve_paragraph_alignment(text)
    para.paragraph_format.line_spacing = profile.docx_line_spacing
    para.paragraph_format.space_after = Pt(profile.docx_paragraph_space_after_pt)

    if preamble:
        para.paragraph_format.left_indent = Mm(0)
        para.paragraph_format.first_line_indent = Mm(0)
    elif _is_structured_line(text):
        para.paragraph_format.left_indent = Mm(profile.docx_list_left_indent_cm * 10)
        para.paragraph_format.first_line_indent = Mm(-profile.docx_list_hanging_cm * 10)
    else:
        para.paragraph_format.left_indent = Mm(0)
        para.paragraph_format.first_line_indent = Mm(profile.docx_first_line_indent_cm * 10)

    run = para.add_run(text)
    _set_run_font(run, profile.body_font_size_pt, font_name)


def _write_heading_paragraph(
    doc: DocxDocument,
    text: str,
    profile: ThaiLegalLayoutProfile,
    font_name: str,
    *,
    title: bool = False,
) -> None:
    para = _new_paragraph(doc)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 42 else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = profile.docx_line_spacing
    para.paragraph_format.space_before = Pt(0 if title else profile.docx_heading_space_before_pt)
    para.paragraph_format.space_after = Pt(
        profile.docx_title_space_after_pt if title else profile.docx_heading_space_after_pt
    )

    run = para.add_run(text)
    _set_run_font(
        run,
        profile.title_font_size_pt if title else profile.heading_font_size_pt,
        font_name,
        bold=True,
    )


def _write_centered_meta(
    doc: DocxDocument,
    text: str,
    profile: ThaiLegalLayoutProfile,
    font_name: str,
) -> None:
    para = _new_paragraph(doc)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(profile.docx_date_space_after_pt)
    run = para.add_run(text)
    _set_run_font(run, profile.date_font_size_pt, font_name)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _write_cell_text(cell, text: str, profile: ThaiLegalLayoutProfile, font_name: str) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = profile.docx_line_spacing
    run = para.add_run(text)
    _set_run_font(run, profile.body_font_size_pt, font_name)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _build_signature_table(
    doc: DocxDocument,
    profile: ThaiLegalLayoutProfile,
    font_name: str,
    left_label: str,
    left_role: str,
    right_label: str,
    right_role: str,
) -> None:
    usable_width = profile.content_width_mm
    spacer_width = usable_width * profile.signature_spacer_ratio
    signer_width = (usable_width - spacer_width) / 2

    table = doc.add_table(rows=3, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borderless(table)

    for row in table.rows:
        _set_cell_width(row.cells[0], signer_width)
        _set_cell_width(row.cells[1], spacer_width)
        _set_cell_width(row.cells[2], signer_width)

    _write_cell_text(table.cell(0, 0), "ลงชื่อ ................................................", profile, font_name)
    _write_cell_text(table.cell(1, 0), f"( {left_label} )", profile, font_name)
    _write_cell_text(table.cell(2, 0), left_role, profile, font_name)

    _write_cell_text(table.cell(0, 1), "", profile, font_name)
    _write_cell_text(table.cell(1, 1), "", profile, font_name)
    _write_cell_text(table.cell(2, 1), "", profile, font_name)

    _write_cell_text(table.cell(0, 2), "ลงชื่อ ................................................", profile, font_name)
    _write_cell_text(table.cell(1, 2), f"( {right_label} )", profile, font_name)
    _write_cell_text(table.cell(2, 2), right_role, profile, font_name)


def generate_contract_docx(
    title: str,
    preamble: str,
    articles: list[ContractArticle],
    parties: list[PartyInfo],
    effective_date: Optional[str] = None,
) -> bytes:
    """
    Generate a Word (.docx) contract document.
    Returns the DOCX file as bytes.
    """

    profile = get_default_thai_legal_layout_profile()
    font_name = _resolve_font_name(profile)
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Mm(profile.margin_top_mm)
        section.bottom_margin = Mm(profile.margin_bottom_mm)
        section.left_margin = Mm(profile.margin_left_mm)
        section.right_margin = Mm(profile.margin_right_mm)

    _write_heading_paragraph(doc, title, profile, font_name, title=True)

    if effective_date:
        _write_centered_meta(doc, f"ลงวันที่ {effective_date}", profile, font_name)

    if preamble:
        for para_text in _iter_paragraphs(preamble):
            _write_body_paragraph(doc, para_text, profile, font_name, preamble=True)

        _new_paragraph(doc)

    for article in articles:
        heading_text = f"ข้อ {article.article_number}  {article.title_th}"
        _write_heading_paragraph(doc, heading_text, profile, font_name)

        for para_text in _iter_paragraphs(article.body_th):
            _write_body_paragraph(doc, para_text, profile, font_name)

    _new_paragraph(doc)
    _new_paragraph(doc)

    if len(parties) >= 2:
        _build_signature_table(
            doc,
            profile,
            font_name,
            _party_display_name(parties[0]),
            _party_role_label(parties[0]),
            _party_display_name(parties[1]),
            _party_role_label(parties[1]),
        )
    elif len(parties) == 1:
        _build_signature_table(
            doc,
            profile,
            font_name,
            _party_display_name(parties[0]),
            _party_role_label(parties[0]),
            "........................................",
            "คู่สัญญา",
        )

    _new_paragraph(doc)
    _new_paragraph(doc)

    _build_signature_table(
        doc,
        profile,
        font_name,
        "........................................",
        "พยาน",
        "........................................",
        "พยาน",
    )

    _new_paragraph(doc)
    footer = _new_paragraph(doc)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    footer_run = footer.add_run("Generated by NeoEM AI Contract Drafting Service")
    _set_run_font(footer_run, profile.footer_font_size_pt, font_name)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
